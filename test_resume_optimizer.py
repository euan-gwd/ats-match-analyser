"""Unit tests for resume_optimizer.py"""

import pytest
import io
from docx import Document
from resume_optimizer import (
    _split_header_and_body,
    _extract_sections,
    build_optimized_resume,
    build_optimized_resume_docx,
)
from ats_scoring import ScoreBreakdown


class TestSplitHeaderAndBody:
    def test_basic_split(self):
        resume = """John Doe
john@example.com
123-456-7890

Experience
Software Engineer at Company A"""
        header, body = _split_header_and_body(resume)
        assert "John Doe" in header
        assert "john@example.com" in header
        assert "Experience" in body

    def test_no_blank_line(self):
        resume = "John Doe\njohn@example.com\nExperience"
        header, body = _split_header_and_body(resume)
        # All goes to header if no blank line
        assert "John Doe" in header

    def test_empty_resume(self):
        header, body = _split_header_and_body("")
        assert header == ""
        assert body == ""


class TestExtractSections:
    def test_extract_experience_section(self):
        resume = """
        John Doe

        Experience
        Software Engineer at Company A
        Developed Python applications

        Education
        BS Computer Science
        """
        sections = _extract_sections(resume)
        assert any("experience" in key.lower() for key in sections.keys())
        assert any("education" in key.lower() for key in sections.keys())

    def test_extract_skills_section(self):
        resume = """
        Skills
        Python, Java, Docker

        Experience
        Software Engineer
        """
        sections = _extract_sections(resume)
        assert any("skills" in key.lower() for key in sections.keys())
        skills_content = next(v for k, v in sections.items() if "skills" in k.lower())
        assert "Python" in skills_content

    def test_case_insensitive_matching(self):
        resume = """
        EXPERIENCE
        Developer role

        EDUCATION
        BS Degree
        """
        sections = _extract_sections(resume)
        assert len(sections) >= 2

    def test_no_sections(self):
        resume = "Just some random text without any section headers"
        sections = _extract_sections(resume)
        # Should return empty dict or only non-standard sections
        assert isinstance(sections, dict)

    def test_multiple_experience_variations(self):
        resume = """
        Professional Experience
        Job 1

        Work Experience
        Job 2
        """
        sections = _extract_sections(resume)
        # Should capture at least one experience section
        experience_sections = [k for k in sections.keys() if "experience" in k.lower()]
        assert len(experience_sections) >= 1


class TestBuildOptimizedResume:
    def test_basic_optimization(self):
        jd = "Python developer with Docker experience"
        resume = """John Doe
john@example.com

Experience
Python Developer at Company A
Used Docker for deployments

Skills
Python, Docker, AWS"""

        breakdown = ScoreBreakdown(
            overall=75.0,
            keyword_similarity=80.0,
            skills_coverage=70.0,
            seniority_alignment=75.0,
            ats_friendliness=80.0,
            posting_recency_factor=100.0,
            missing_keywords=["kubernetes"],
            matched_keywords=["python", "docker"],
            seniority_explanation="Good match",
            ats_reasons=[],
            recency_explanation="Fresh posting",
        )

        optimized = build_optimized_resume(jd, resume, breakdown)

        assert isinstance(optimized, str)
        assert len(optimized) > 0
        assert "John Doe" in optimized
        assert "john@example.com" in optimized

    def test_preserves_contact_info(self):
        resume = """Jane Smith
jane@example.com
555-1234

Experience
Developer"""

        breakdown = ScoreBreakdown(
            overall=60.0,
            keyword_similarity=60.0,
            skills_coverage=60.0,
            seniority_alignment=60.0,
            ats_friendliness=60.0,
            posting_recency_factor=100.0,
            missing_keywords=[],
            matched_keywords=["python"],
            seniority_explanation="",
            ats_reasons=[],
            recency_explanation="",
        )

        optimized = build_optimized_resume("", resume, breakdown)
        assert "Jane Smith" in optimized
        assert "jane@example.com" in optimized

    def test_includes_matched_keywords(self):
        resume = "Developer with Python and Docker skills"
        breakdown = ScoreBreakdown(
            overall=70.0,
            keyword_similarity=70.0,
            skills_coverage=70.0,
            seniority_alignment=70.0,
            ats_friendliness=70.0,
            posting_recency_factor=100.0,
            missing_keywords=[],
            matched_keywords=["python", "docker", "aws"],
            seniority_explanation="",
            ats_reasons=[],
            recency_explanation="",
        )

        optimized = build_optimized_resume("", resume, breakdown)
        assert "python" in optimized.lower()
        assert "docker" in optimized.lower()

    def test_empty_resume(self):
        breakdown = ScoreBreakdown(
            overall=0.0,
            keyword_similarity=0.0,
            skills_coverage=0.0,
            seniority_alignment=0.0,
            ats_friendliness=0.0,
            posting_recency_factor=100.0,
            missing_keywords=[],
            matched_keywords=[],
            seniority_explanation="",
            ats_reasons=[],
            recency_explanation="",
        )

        optimized = build_optimized_resume("", "", breakdown)
        assert isinstance(optimized, str)


class TestBuildOptimizedResumeDocx:
    def test_creates_valid_docx(self):
        resume = """John Doe
john@example.com

Experience
Software Engineer at Company A

Skills
Python, Docker"""

        breakdown = ScoreBreakdown(
            overall=75.0,
            keyword_similarity=80.0,
            skills_coverage=70.0,
            seniority_alignment=75.0,
            ats_friendliness=80.0,
            posting_recency_factor=100.0,
            missing_keywords=[],
            matched_keywords=["python", "docker"],
            seniority_explanation="Good match",
            ats_reasons=[],
            recency_explanation="Fresh posting",
        )

        docx_buffer = build_optimized_resume_docx("", resume, breakdown)

        assert isinstance(docx_buffer, io.BytesIO)
        assert docx_buffer.tell() == 0  # Should be at start

        # Verify it's a valid docx
        doc = Document(docx_buffer)
        assert len(doc.paragraphs) > 0

    def test_docx_contains_name(self):
        resume = """Jane Smith
jane@example.com

Experience
Developer"""

        breakdown = ScoreBreakdown(
            overall=60.0,
            keyword_similarity=60.0,
            skills_coverage=60.0,
            seniority_alignment=60.0,
            ats_friendliness=60.0,
            posting_recency_factor=100.0,
            missing_keywords=[],
            matched_keywords=["python"],
            seniority_explanation="",
            ats_reasons=[],
            recency_explanation="",
        )

        docx_buffer = build_optimized_resume_docx("", resume, breakdown)
        doc = Document(docx_buffer)

        text_content = "\n".join([p.text for p in doc.paragraphs])
        assert "Jane Smith" in text_content

    def test_docx_has_sections(self):
        resume = """John Doe
john@example.com

Experience
Software Engineer

Education
BS Computer Science

Skills
Python, Docker"""

        breakdown = ScoreBreakdown(
            overall=75.0,
            keyword_similarity=75.0,
            skills_coverage=75.0,
            seniority_alignment=75.0,
            ats_friendliness=75.0,
            posting_recency_factor=100.0,
            missing_keywords=[],
            matched_keywords=["python", "docker"],
            seniority_explanation="",
            ats_reasons=[],
            recency_explanation="",
        )

        docx_buffer = build_optimized_resume_docx("", resume, breakdown)
        doc = Document(docx_buffer)

        text_content = "\n".join([p.text for p in doc.paragraphs])
        # Should have standard section headers
        assert any(header in text_content.upper() for header in ["EXPERIENCE", "EDUCATION", "SKILLS"])

    def test_docx_includes_keywords(self):
        resume = "Developer with Python skills"
        breakdown = ScoreBreakdown(
            overall=70.0,
            keyword_similarity=70.0,
            skills_coverage=70.0,
            seniority_alignment=70.0,
            ats_friendliness=70.0,
            posting_recency_factor=100.0,
            missing_keywords=[],
            matched_keywords=["python", "docker", "aws", "kubernetes"],
            seniority_explanation="",
            ats_reasons=[],
            recency_explanation="",
        )

        docx_buffer = build_optimized_resume_docx("", resume, breakdown)
        doc = Document(docx_buffer)

        text_content = "\n".join([p.text for p in doc.paragraphs])
        assert "python" in text_content.lower()

    def test_empty_resume_docx(self):
        breakdown = ScoreBreakdown(
            overall=0.0,
            keyword_similarity=0.0,
            skills_coverage=0.0,
            seniority_alignment=0.0,
            ats_friendliness=0.0,
            posting_recency_factor=100.0,
            missing_keywords=[],
            matched_keywords=[],
            seniority_explanation="",
            ats_reasons=[],
            recency_explanation="",
        )

        docx_buffer = build_optimized_resume_docx("", "", breakdown)
        assert isinstance(docx_buffer, io.BytesIO)

        # Should still create a valid (though empty) document
        doc = Document(docx_buffer)
        assert doc is not None
        assert hasattr(doc, 'paragraphs')

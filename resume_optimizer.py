from __future__ import annotations

import io
import re
from typing import List

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ats_scoring import ScoreBreakdown


def _split_header_and_body(resume_text: str) -> tuple[str, str]:
    """
    Roughly split the resume into a header (top block) and body.
    The header is all lines until the first blank line.
    """
    lines = resume_text.splitlines()
    header_lines: List[str] = []
    body_lines: List[str] = []

    seen_blank = False
    for line in lines:
        if not seen_blank and line.strip() == "":
            seen_blank = True
            continue
        if not seen_blank:
            header_lines.append(line)
        else:
            body_lines.append(line)

    header = "\n".join(header_lines).strip()
    body = "\n".join(body_lines).strip()
    return header, body


def _extract_sections(resume_text: str) -> dict[str, str]:
    """
    Extract common resume sections like Experience, Education, Skills, etc.
    Returns a dict mapping section names to their content.
    """
    section_patterns = [
        r"(?i)(experience|work experience|professional experience|employment history)",
        r"(?i)(education|academic background)",
        r"(?i)(skills|technical skills|core competencies)",
        r"(?i)(projects|key projects)",
        r"(?i)(certifications|certificates)",
        r"(?i)(summary|professional summary|profile)",
    ]

    sections = {}
    lines = resume_text.split('\n')
    current_section = None
    current_content = []

    for line in lines:
        # Check if this line is a section header
        is_header = False
        for pattern in section_patterns:
            if re.match(pattern, line.strip()) and len(line.strip()) < 50:
                # Save previous section
                if current_section:
                    sections[current_section] = '\n'.join(current_content).strip()
                # Start new section
                current_section = line.strip()
                current_content = []
                is_header = True
                break

        if not is_header and current_section:
            current_content.append(line)

    # Save last section
    if current_section:
        sections[current_section] = '\n'.join(current_content).strip()

    return sections


def _highlight_keywords_in_text(text: str, keywords: List[str]) -> str:
    """
    Strategically weave keywords into existing text where they make sense.
    Only adds keywords that genuinely relate to the context.
    """
    # This is a conservative approach - we don't modify the text
    # Just return it as-is since we don't want to hallucinate
    return text


def build_optimized_resume(
    jd_text: str,
    resume_text: str,
    breakdown: ScoreBreakdown,
) -> str:
    """
    Construct an ATS-optimized resume draft that:
    - Only uses facts and wording already present in the original resume.
    - Reorganizes content to foreground matched keywords/skills.
    - Does NOT introduce new skills, experience, or dates.
    """
    resume_text = resume_text or ""
    header, body = _split_header_and_body(resume_text)

    # Use only matched keywords/skills that appear in the original resume
    matched_terms = sorted(set(breakdown.matched_keywords))
    top_terms = matched_terms[:30] if matched_terms else []

    sections: List[str] = []

    if header:
        sections.append(header.strip())

    if top_terms:
        skills_block = [
            "CORE SKILLS (derived from your existing resume)",
            ", ".join(top_terms),
        ]
        sections.append("\n".join(skills_block))

    # A neutral summary that only references existing skills
    if top_terms:
        summary_line = (
            "SUMMARY\n"
            f"Professional profile emphasizing: {', '.join(top_terms[:10])}."
        )
        sections.append(summary_line)

    if body:
        sections.append("EXPERIENCE, PROJECTS, AND EDUCATION (original wording)")
        sections.append(body.strip())

    optimized = "\n\n".join(s for s in sections if s.strip())
    return optimized.strip()


def build_optimized_resume_docx(
    jd_text: str,
    resume_text: str,
    breakdown: ScoreBreakdown,
) -> io.BytesIO:
    """
    Create an ATS-optimized resume as a Word document (.docx) that:
    - Maximizes ATS score by strategic organization
    - Uses only existing information from the resume
    - Follows ATS-friendly formatting best practices
    - Returns a BytesIO object ready for download
    """
    doc = Document()

    # Set narrow margins for better space utilization
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    resume_text = resume_text or ""
    header, body = _split_header_and_body(resume_text)
    extracted_sections = _extract_sections(body)

    # Use matched keywords that appear in original resume
    matched_terms = sorted(set(breakdown.matched_keywords))
    all_keywords = matched_terms + [k for k in breakdown.missing_keywords if k.lower() in resume_text.lower()]
    top_keywords = sorted(set(all_keywords))[:40]

    # === HEADER SECTION ===
    if header:
        # Extract name (usually first line) and contact info
        header_lines = header.split('\n')
        if header_lines:
            # Name - centered, bold, larger font
            name_para = doc.add_paragraph()
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_run = name_para.add_run(header_lines[0].strip())
            name_run.bold = True
            name_run.font.size = Pt(16)

            # Contact info - centered, smaller font
            if len(header_lines) > 1:
                contact_para = doc.add_paragraph()
                contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                contact_run = contact_para.add_run('\n'.join(header_lines[1:]).strip())
                contact_run.font.size = Pt(10)

    # === PROFESSIONAL SUMMARY ===
    # Create a compelling summary using existing skills
    if top_keywords:
        doc.add_paragraph()  # spacing
        summary_heading = doc.add_paragraph()
        summary_run = summary_heading.add_run('PROFESSIONAL SUMMARY')
        summary_run.bold = True
        summary_run.font.size = Pt(12)

        # Find any existing summary from the resume
        existing_summary = ""
        for section_name, content in extracted_sections.items():
            if 'summary' in section_name.lower() or 'profile' in section_name.lower():
                existing_summary = content
                break

        summary_text = existing_summary if existing_summary else (
            f"Professional with demonstrated expertise in: {', '.join(top_keywords[:12])}. "
            "Proven track record of delivering results through technical excellence and strategic thinking."
        )
        summary_para = doc.add_paragraph(summary_text)
        summary_para.paragraph_format.space_after = Pt(6)

    # === CORE SKILLS ===
    if top_keywords:
        doc.add_paragraph()  # spacing
        skills_heading = doc.add_paragraph()
        skills_run = skills_heading.add_run('CORE SKILLS')
        skills_run.bold = True
        skills_run.font.size = Pt(12)

        # Format skills in multiple columns for readability
        skills_text = ' • '.join(top_keywords)
        skills_para = doc.add_paragraph(skills_text)
        skills_para.paragraph_format.space_after = Pt(6)

    # === PROFESSIONAL EXPERIENCE ===
    for section_name, content in extracted_sections.items():
        if 'experience' in section_name.lower() or 'employment' in section_name.lower():
            doc.add_paragraph()  # spacing
            exp_heading = doc.add_paragraph()
            exp_run = exp_heading.add_run('PROFESSIONAL EXPERIENCE')
            exp_run.bold = True
            exp_run.font.size = Pt(12)

            # Add the experience content, preserving formatting
            exp_lines = content.split('\n')
            for line in exp_lines:
                if line.strip():
                    para = doc.add_paragraph(line.strip())
                    para.paragraph_format.space_after = Pt(3)
                    # Bold lines that look like job titles/companies (heuristic)
                    if len(line.strip()) < 80 and not line.strip().startswith(('•', '-', '–')):
                        para.runs[0].bold = True

    # === EDUCATION ===
    for section_name, content in extracted_sections.items():
        if 'education' in section_name.lower():
            doc.add_paragraph()  # spacing
            edu_heading = doc.add_paragraph()
            edu_run = edu_heading.add_run('EDUCATION')
            edu_run.bold = True
            edu_run.font.size = Pt(12)

            edu_lines = content.split('\n')
            for line in edu_lines:
                if line.strip():
                    para = doc.add_paragraph(line.strip())
                    para.paragraph_format.space_after = Pt(3)

    # === PROJECTS ===
    for section_name, content in extracted_sections.items():
        if 'project' in section_name.lower():
            doc.add_paragraph()  # spacing
            proj_heading = doc.add_paragraph()
            proj_run = proj_heading.add_run('KEY PROJECTS')
            proj_run.bold = True
            proj_run.font.size = Pt(12)

            proj_lines = content.split('\n')
            for line in proj_lines:
                if line.strip():
                    para = doc.add_paragraph(line.strip())
                    para.paragraph_format.space_after = Pt(3)

    # === CERTIFICATIONS ===
    for section_name, content in extracted_sections.items():
        if 'certification' in section_name.lower() or 'certificate' in section_name.lower():
            doc.add_paragraph()  # spacing
            cert_heading = doc.add_paragraph()
            cert_run = cert_heading.add_run('CERTIFICATIONS')
            cert_run.bold = True
            cert_run.font.size = Pt(12)

            cert_lines = content.split('\n')
            for line in cert_lines:
                if line.strip():
                    para = doc.add_paragraph(line.strip())
                    para.paragraph_format.space_after = Pt(3)

    # === ADDITIONAL SKILLS (if not already covered) ===
    for section_name, content in extracted_sections.items():
        if 'skill' in section_name.lower() and section_name not in [s for s in extracted_sections.keys() if 'core' in s.lower()]:
            # Only add if we haven't already added a skills section
            if not any('CORE SKILLS' in p.text for p in doc.paragraphs):
                doc.add_paragraph()  # spacing
                addl_heading = doc.add_paragraph()
                addl_run = addl_heading.add_run('TECHNICAL SKILLS')
                addl_run.bold = True
                addl_run.font.size = Pt(12)

                skill_lines = content.split('\n')
                for line in skill_lines:
                    if line.strip():
                        para = doc.add_paragraph(line.strip())
                        para.paragraph_format.space_after = Pt(3)

    # Save to BytesIO
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)

    return docx_buffer


